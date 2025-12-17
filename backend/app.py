import os
import io
from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from datetime import datetime
from cryptography.fernet import Fernet
from werkzeug.utils import secure_filename
from openai import OpenAI
from docx import Document
from flask_socketio import SocketIO, emit, join_room, leave_room

app = Flask(__name__)
# کلید سکرت برنامه (بهتر است از متغیر محیطی خوانده شود)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'default-secret-key')

CORS(app, resources={r"/*": {"origins": "*"}})
# تنظیم سوکت برای جلوگیری از تداخل با eventlet
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# تنظیمات دیتابیس
db_uri = os.environ.get('SQLALCHEMY_DATABASE_URI', 'postgresql://myuser:mypassword@127.0.0.1:5440/meeting_db')
app.config['SQLALCHEMY_DATABASE_URI'] = db_uri
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# --- تنظیمات امنیتی کلیدها ---

# 1. کلید رمزنگاری دیتابیس (Fernet)
# اگر در محیط سرور (ENV) نبود، یک کلید پیش‌فرض استفاده کن (فقط برای تست)
# نکته: در سرور واقعی باید این را در Environment Variables ست کنید
default_fernet_key = b'gJ5v9KzX7V4lM2nB8cQ1wE3rT6yU0iO9pA8sD7fG6hJ=' 
encryption_key_str = os.environ.get('FERNET_KEY')
if encryption_key_str:
    ENCRYPTION_KEY = encryption_key_str.encode()
else:
    ENCRYPTION_KEY = default_fernet_key

cipher_suite = Fernet(ENCRYPTION_KEY)

# 2. کلید OpenAI
# کلید را از متغیر محیطی می‌خواند. اگر نبود، ارور نمی‌دهد اما کار نمی‌کند.
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY", "YOUR_API_KEY_HERE"))

db = SQLAlchemy(app)

# --- مدل‌ها ---
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    full_name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    phone = db.Column(db.String(20))
    company = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    meets = db.relationship('Meet', backref='user_info', lazy=True)

class MeetingRoom(db.Model):
    __tablename__ = 'meeting_rooms'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    description = db.Column(db.String(200))

class Meet(db.Model):
    __tablename__ = 'meet'
    meet_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    gmail = db.Column(db.String(100))
    session_id = db.Column(db.Integer, db.ForeignKey('meeting_rooms.id'), nullable=False)
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)

class RecordText(db.Model):
    __tablename__ = 'record_text'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    meet_id = db.Column(db.Integer, db.ForeignKey('meet.meet_id'), nullable=False)
    text = db.Column(db.LargeBinary, nullable=False)
    audio_blob = db.Column(db.LargeBinary, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# --- روت‌ها ---
@app.route('/register', methods=['POST'])
def register():
    data = request.json
    if User.query.filter_by(email=data['email']).first(): return jsonify({'message': 'ایمیل تکراری است'}), 409
    try:
        new_user = User(full_name=data['full_name'], email=data['email'], phone=data.get('phone'), company=data.get('company'))
        db.session.add(new_user)
        db.session.commit()
        return jsonify({'message': 'ثبت نام موفق', 'user_id': new_user.id, 'user_name': new_user.full_name, 'email': new_user.email}), 201
    except Exception as e: return jsonify({'message': str(e)}), 500

@app.route('/login', methods=['POST'])
def login():
    data = request.json
    user = User.query.filter_by(email=data['email']).first()
    if user: return jsonify({'message': 'ورود موفق', 'user_id': user.id, 'user_name': user.full_name, 'email': user.email}), 200
    return jsonify({'message': 'کاربر یافت نشد'}), 404

@app.route('/meetings', methods=['GET'])
def get_meetings():
    rooms = MeetingRoom.query.all()
    return jsonify([{'id': r.id, 'title': r.title, 'description': r.description} for r in rooms])

@app.route('/join_meeting', methods=['POST'])
def join_meeting():
    data = request.json
    try:
        new_meet = Meet(user_id=data['user_id'], gmail=data['email'], session_id=data['session_id'])
        db.session.add(new_meet)
        db.session.commit()
        return jsonify({'message': 'وارد جلسه شدید', 'meet_id': new_meet.meet_id}), 200
    except Exception as e: return jsonify({'message': str(e)}), 500

@app.route('/meeting_members/<int:session_id>', methods=['GET'])
def get_meeting_members(session_id):
    results = db.session.query(Meet, User).join(User, Meet.user_id == User.id).filter(Meet.session_id == session_id).all()
    seen_users = set()
    result = []
    for meet, user in results:
        if user.id not in seen_users:
            result.append({'user_name': user.full_name, 'email': user.email})
            seen_users.add(user.id)
    return jsonify(result)

@app.route('/upload_audio', methods=['POST'])
def upload_audio():
    print("\n\n========================================", flush=True)
    print("--- درخواست آپلود صدا دریافت شد ---", flush=True)
    
    if 'file' not in request.files: 
        print("❌ خطا: فایل در درخواست نیست", flush=True)
        return jsonify({'message': 'فایل نیست'}), 400
    
    file = request.files['file']
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)
    print(f"📁 سایز فایل دریافتی: {file_size} بایت", flush=True)

    if file_size < 1000:
        print("⚠️ هشدار: فایل صوتی خیلی کوچک یا خالی است!", flush=True)

    user_id = request.form.get('user_id')
    meet_id = request.form.get('meet_id')

    try:
        filename = secure_filename(f"temp_{user_id}_{datetime.now().timestamp()}.webm")
        file.save(filename)
        
        with open(filename, "rb") as f: audio_data = f.read()
        encrypted_audio = cipher_suite.encrypt(audio_data)

        print("⏳ در حال ارسال به OpenAI...", flush=True)
        try:
            with open(filename, "rb") as audio_file:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1", 
                    file=audio_file, 
                    language="fa"
                )
            extracted_text = transcript.text
            print(f"✅ متن استخراج شده: {extracted_text}", flush=True)
        except Exception as ai_error:
            print(f"❌ خطای OpenAI: {ai_error}", flush=True)
            extracted_text = f"خطا: {str(ai_error)}"

        encrypted_text = cipher_suite.encrypt(extracted_text.encode('utf-8'))

        new_record = RecordText(user_id=user_id, meet_id=meet_id, text=encrypted_text, audio_blob=encrypted_audio)
        db.session.add(new_record)
        db.session.commit()
        
        if os.path.exists(filename): os.remove(filename)
            
        return jsonify({'message': 'ذخیره شد', 'text_preview': extracted_text}), 200
    except Exception as e:
        print(f"❌ Error: {e}", flush=True)
        return jsonify({'message': str(e)}), 500

@app.route('/get_session_records/<int:session_id>', methods=['GET'])
def get_session_records(session_id):
    my_meet_id = request.args.get('my_meet_id')
    if not my_meet_id: return jsonify({'message': 'شناسه ورود کاربر مشخص نیست'}), 400

    current_user_meet = Meet.query.get(my_meet_id)
    if not current_user_meet: return jsonify({'message': 'اطلاعات ورود یافت نشد'}), 404
        
    user_join_time = current_user_meet.joined_at

    results = db.session.query(RecordText, User)\
        .join(Meet, RecordText.meet_id == Meet.meet_id)\
        .join(User, RecordText.user_id == User.id)\
        .filter(Meet.session_id == session_id)\
        .filter(RecordText.created_at >= user_join_time) \
        .order_by(RecordText.created_at.asc())\
        .all()
    
    output = []
    for record, user in results:
        try:
            decrypted_text = cipher_suite.decrypt(record.text).decode('utf-8')
            output.append({'id': record.id, 'user_name': user.full_name, 'text': decrypted_text, 'time': record.created_at.strftime("%H:%M:%S")})
        except: output.append({'id': record.id, 'user_name': user.full_name, 'text': '(خطا در رمزگشایی)', 'time': ''})
    return jsonify(output)

@app.route('/download_word/<int:session_id>', methods=['GET'])
def download_word(session_id):
    results = db.session.query(RecordText, User).join(Meet, RecordText.meet_id == Meet.meet_id).join(User, RecordText.user_id == User.id).filter(Meet.session_id == session_id).order_by(RecordText.created_at.asc()).all()
    document = Document()
    document.add_heading('گزارش کامل جلسه', 0)
    for record, user in results:
        try:
            text = cipher_suite.decrypt(record.text).decode('utf-8')
            time = record.created_at.strftime("%Y-%m-%d %H:%M:%S")
            p = document.add_paragraph()
            runner = p.add_run(f"[{time}] {user.full_name}: ")
            runner.bold = True
            p.add_run(text)
        except: continue
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True, download_name=f'report_session_{session_id}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# --- سوکت ---
@socketio.on('join')
def on_join(data):
    room = data['session_id']
    join_room(room)
    print(f"User joined room: {room}")

@socketio.on('leave')
def on_leave(data):
    room = data['session_id']
    leave_room(room)

@socketio.on('voice_stream')
def handle_voice_stream(data):
    room = data['session_id']
    emit('play_audio', data, to=room, include_self=False)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        if not MeetingRoom.query.first():
            db.session.add(MeetingRoom(title="جلسه هیئت مدیره", description="بررسی عملکرد سالانه"))
            db.session.add(MeetingRoom(title="تیم فنی (Tech)", description="بررسی باگ‌های پروژه"))
            db.session.add(MeetingRoom(title="جلسه بازاریابی", description="کمپین نوروزی"))
            db.session.commit()
            
    # تنظیمات اجرا برای سرور (بدون SSL داخلی، چون Nginx هندل می‌کند)
    # اگر روی لوکال تست می‌کنید و HTTPS می‌خواهید، می‌توانید ssl_context را برگردانید
    # اما برای گیت و سرور واقعی، این حالت استاندارد است:
    socketio.run(app, host='0.0.0.0', port=5000, debug=False)